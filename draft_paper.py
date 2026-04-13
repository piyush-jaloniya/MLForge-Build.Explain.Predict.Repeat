from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_draft():
    doc = Document()
    
    # Title
    doc.add_heading('MLForge: Build. Explain. Predict. Repeat.', 0)
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "Machine learning is widely used, but building reliable models usually requires a strong background in software engineering and statistics. While automated machine learning (AutoML) tools help lower this barrier, they often create a new problem: they act as black boxes. Users put data in and get predictions out, but they rarely get to see exactly how those predictions were made. To address this, we built a self-hostable AutoML web application designed around transparency and ease of use. The platform handles data preprocessing, hyperparameter tuning via Optuna, and trains models using ten common algorithms entirely within the browser interface. Beyond just automating tasks, we focused on helping the user understand the process. We implemented a reinforcement learning advisory system using Thompson Sampling to suggest models based on how they actually perform during cross-validation. We also integrated Google Gemini to provide plain-text explanations of charts and metrics. Because LLMs can hallucinate or slow down the system, we strictly limited the Gemini context budget to under 2KB, passing only scalar metrics rather than raw data. Finally, we embedded Explainable AI (XAI) visualizations natively into the dashboard, so users can verify model logic globally and locally. Ultimately, this platform demonstrates that removing the need for code doesn't have to mean hiding the underlying math."
    )
    
    # Section 1
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "For a domain expert like a biologist or a financial analyst, setting up a machine learning pipeline from scratch can be frustrating. Between dealing with missing values, properly encoding categories, and avoiding data leakage, the technical overhead is high. This learning curve has driven the popularity of AutoML platforms, which are built to reduce human intervention. However, in practice, fully automating a pipeline can actually degrade user trust. When the tool hides the algorithm selection process and abstracts away feature engineering, users are forced to accept the final predictions on blind faith."
    )
    doc.add_paragraph(
        "This project started with a specific goal: to build MLForge as a system that automates the tedious parts of data science while explicitly showing its work. We realized that existing no-code platforms usually optimize for speed and ease, often treating interpretability as an afterthought. We wanted to see what happens when explainability and AI guidance are treated as core features of the user interface, rather than optional add-ons."
    )
    doc.add_paragraph(
        "Developing this system required several practical tradeoffs. For example, to keep the UI responsive during interactive workflows, we chose to manage dataset instances and model training artifacts entirely in-memory using a FastAPI backend. While this avoids the latency of querying an external database, it forced us to implement strict LRU (Least Recently Used) eviction policies to keep memory consumption stable. With that architecture in place, we focused our contributions locally on three areas:"
    )
    doc.add_paragraph(
        "First, we integrated established explainability libraries directly into the training loop. We calculate SHAP (Shapley Additive exPlanations) values and permutation importance asynchronously in the background. This means users don't have to write custom scripts to generate beeswarm or waterfall plots; the visualizations compile automatically as part of the model evaluation."
    )
    doc.add_paragraph(
        "Second, we built an active advisory framework for model selection. Choosing the right algorithm often comes down to trial and error. We framed this as a multi-armed bandit problem, using a Thompson Sampling module to track empirical performance over time. The system learns from previous sessions and gradually updates its recommendations, guiding users toward historically successful algorithms for similar data."
    )
    doc.add_paragraph(
        "Finally, we integrated generative AI (Google Gemini) to serve as an embedded assistant. Rather than letting the LLM write code or run the platform—which introduces significant safety risks—we restricted it to an advisory role. By feeding the LLM tightly constrained context windows composed only of numeric metrics and schema data, the assistant translates complex evaluation metrics into readable summaries without accessing raw user data."
    )
    doc.add_paragraph(
        "The rest of this paper explores the design and implementation of this system. Section 2 reviews related work in AutoML and algorithm selection. Section 3 outlines the React and FastAPI architecture, followed by a breakdown of the ML pipeline in Section 4. We then detail the XAI integration (Section 5), hyperparameter optimization (Section 6), and the reinforcement learning advisor (Section 7). After discussing the generative AI constraints (Section 8) and visualization engine (Section 9), we cover specific engineering challenges we encountered during development in Section 10. The paper concludes with performance evaluations in Section 11 and directions for future work."
    )

    # Section 2
    doc.add_heading('2. Related Work', level=1)
    
    doc.add_heading('2.1 No-Code and Low-Code ML Platforms', level=2)
    doc.add_paragraph(
        "There's no shortage of tools aimed at simplifying machine learning. Enterprise solutions like Google Cloud AutoML and DataRobot are incredibly robust, handling everything from data ingestion to Kubernetes deployment. The catch is that they are expensive and heavily tied to proprietary cloud infrastructure, making them inaccessible for many students or small clinics. On the other end of the spectrum are Python libraries like PyCaret and Auto-sklearn. These are excellent for reducing boilerplate code, but they still assume the user knows how to set up a Python environment and read a traceback error. We designed our platform to sit between these extremes. It offers the visual ease of an enterprise web app, but it's entirely open-source and can be self-hosted locally on a laptop."
    )
    
    doc.add_heading('2.2 Explainable AI (XAI) Tools', level=2)
    doc.add_paragraph(
        "As models grow more complex, the need for post-hoc explanation tools has become urgent. LIME and SHAP are currently the industry standards for extracting local feature importance. SHAP is particularly valued because of its solid foundation in cooperative game theory. However, the main issue with SHAP in a web application is its computational cost—running a TreeExplainer over a large dataset blocks the event loop and freezes the UI. In most existing tools, XAI is treated as a separate, slow diagnostic step. In our platform, we push SHAP calculations to background threads. By capping the background sample size for the explainers, we ensure that both global beeswarm plots and local waterfall plots render alongside standard accuracy metrics without crashing the server."
    )

    doc.add_heading('2.3 Hyperparameter Optimization', level=2)
    doc.add_paragraph(
        "Finding the right parameters for a model like XGBoost used to rely on exhaustive Grid Search or random sampling. These methods are notoriously inefficient because they don't learn from previous trials. Modern frameworks like Optuna and Ray Tune use Bayesian optimization approaches, specifically the Tree-structured Parzen Estimator (TPE), to guess which parameter combinations are actually worth testing. We natively integrated Optuna because of its lack of heavy infrastructure dependencies compared to Ray. Rather than running the tuning process invisibly, the platform streams the trial history and objective scores to the frontend, making the optimization process completely observable."
    )

    doc.add_heading('2.4 Reinforcement Learning for Algorithm Selection', level=2)
    doc.add_paragraph(
        "A common struggle for beginners is knowing whether they should use a Random Forest, a Support Vector Machine, or simple Logistic Regression for a new dataset. This is formally known as the algorithm selection problem. Traditional AutoML pipelines solve this via brute force, running large benchmark tournaments that consume massive compute resources. A leaner approach is to look at model selection as a reinforcement learning task, specifically a multi-armed bandit problem. If each algorithm is an 'arm', the system can balance testing new models (exploration) against recommending models that usually perform well (exploitation). We implemented a simple Beta-Bernoulli bandit with Thompson Sampling for this exact purpose. As the user runs more training jobs, the system updates its internal success priors, creating an advisory engine that gets smarter with usage."
    )

    doc.add_heading('2.5 LLMs in Data Science', level=2)
    doc.add_paragraph(
        "Lately, there has been a massive push to use Large Language Models (LLMs) to write data science scripts. Tools like GitHub Copilot or advanced Jupyter agents can automatically generate Pandas code based on a prompt. However, from a system design perspective, letting an LLM autonomously write and execute code introduces unpredictable failure states. If the LLM misunderstands the schema, the entire pipeline crashes. We took a much more defensive approach. Instead of code generation, we use the Google Gemini API strictly for translation. The frontend parses data quality issues and scalar metrics (like RMSE or F1 scores) into a compact JSON payload. The LLM then reads this numeric data and generates conversational, plain-text explanations for the user. Keeping the context packet under a 2KB budget not only guarantees low latency but also avoids uploading raw, potentially sensitive user data to an external API."
    )

    doc.add_paragraph(
        "Overall, while tools exist for every individual part of the ML lifecycle, there are very few open-source projects that try to knit AutoML, Bayesian optimization, XAI, and LLM assistance into a single, cohesive interface. The rest of this paper focuses on the architectural decisions required to make these separate systems talk to each other smoothly."
    )

    doc.save('mlforge_research_paper.docx')
    print("Created mlforge_research_paper.docx successfully.")

if __name__ == '__main__':
    create_draft()
